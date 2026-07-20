"""
===========================================================
AulaMind Enterprise 3.0
services/export_service.py
-----------------------------------------------------------
Servicio de Exportación
- Microsoft Word (.docx)
- PDF
-----------------------------------------------------------
Biotecno Chile
===========================================================
"""

from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.enums import TA_CENTER


class ExportService:

    # =====================================================
    # WORD
    # =====================================================

    @staticmethod
    def export_word(
        title,
        content,
        teacher=None,
        school=None,
        subject=None,
        course=None
    ):

        document = Document()

        title_style = document.styles["Title"]
        title_style.font.size = Pt(22)

        heading = document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph()

        if school:

            p = document.add_paragraph()
            p.add_run("Establecimiento: ").bold = True
            p.add_run(school)

        if teacher:

            p = document.add_paragraph()
            p.add_run("Profesor(a): ").bold = True
            p.add_run(teacher)

        if subject:

            p = document.add_paragraph()
            p.add_run("Asignatura: ").bold = True
            p.add_run(subject)

        if course:

            p = document.add_paragraph()
            p.add_run("Curso: ").bold = True
            p.add_run(course)

        p = document.add_paragraph()
        p.add_run("Fecha: ").bold = True
        p.add_run(datetime.now().strftime("%d-%m-%Y"))

        document.add_paragraph()

        for linea in content.split("\n"):

            document.add_paragraph(linea)

        stream = BytesIO()

        document.save(stream)

        stream.seek(0)

        return stream

    # =====================================================
    # PDF
    # =====================================================

    @staticmethod
    def export_pdf(
        title,
        content,
        teacher=None,
        school=None,
        subject=None,
        course=None
    ):

        stream = BytesIO()

        doc = SimpleDocTemplate(stream)

        styles = getSampleStyleSheet()

        style_title = styles["Heading1"]
        style_title.alignment = TA_CENTER

        elements = []

        elements.append(
            Paragraph(title, style_title)
        )

        elements.append(
            Paragraph("<br/>", styles["Normal"])
        )

        if school:

            elements.append(
                Paragraph(
                    f"<b>Establecimiento:</b> {school}",
                    styles["Normal"]
                )
            )

        if teacher:

            elements.append(
                Paragraph(
                    f"<b>Profesor:</b> {teacher}",
                    styles["Normal"]
                )
            )

        if subject:

            elements.append(
                Paragraph(
                    f"<b>Asignatura:</b> {subject}",
                    styles["Normal"]
                )
            )

        if course:

            elements.append(
                Paragraph(
                    f"<b>Curso:</b> {course}",
                    styles["Normal"]
                )
            )

        elements.append(
            Paragraph(
                f"<b>Fecha:</b> {datetime.now().strftime('%d-%m-%Y')}",
                styles["Normal"]
            )
        )

        elements.append(
            Paragraph("<br/>", styles["Normal"])
        )

        for linea in content.split("\n"):

            linea = linea.strip()

            if not linea:

                linea = "&nbsp;"

            elements.append(

                Paragraph(
                    linea.replace(" ", "&nbsp;"),
                    styles["BodyText"]
                )

            )

        doc.build(elements)

        stream.seek(0)

        return stream

    # =====================================================
    # NOMBRE DE ARCHIVO
    # =====================================================

    @staticmethod
    def filename(document_type, subject, course):

        subject = (subject or "").replace(" ", "_")

        course = (course or "").replace(" ", "_")

        today = datetime.now().strftime("%Y%m%d")

        return f"{document_type}_{subject}_{course}_{today}"



export_service = ExportService()